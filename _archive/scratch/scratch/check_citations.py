import re
from docx import Document

docx_path = r"C:\Users\user\Desktop\test_html\碩論\碩論vX.docx"
doc = Document(docx_path)

# Extract paragraphs text
paragraphs = [p.text for p in doc.paragraphs]

# 1. Extract citations in text
# Typically: Name et al. (Year) or (Name et al., Year) or Name (Year)
# Regex to find potential academic citations: e.g., Zhang et al. (2020), Devlin et al. (2019), Yang et al. (2016)
citation_patterns = [
    r'[A-Za-z\-]+ et al\.\s*\(\d{4}\)',
    r'\([A-Za-z\-]+ et al\.,\s*\d{4}\)',
    r'[A-Z][a-zA-Z\-]+\s*\(\d{4}\)',
    r'\([A-Z][a-zA-Z\-]+,\s*\d{4}\)'
]

text_citations = set()
for text in paragraphs:
    for pattern in citation_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            text_citations.add(match)

# Get the References paragraph text (usually starts after a heading)
references_text = []
is_ref = False
for p in doc.paragraphs:
    if "參考文獻" in p.text or "References" in p.text:
        is_ref = True
    if is_ref:
        references_text.append(p.text)

ref_block = "\n".join(references_text)

print("--- Citations in Text ---")
for c in sorted(text_citations):
    print(c)

print("\n--- Checking if citations are in Reference list ---")
# Simple check
for cit in text_citations:
    # Extract author name
    name_match = re.search(r'([A-Za-z\-]+)', cit)
    if name_match:
        name = name_match.group(1)
        if name.lower() not in ref_block.lower():
            print(f"WARNING: Citation '{cit}' might be missing from References!")
        else:
            print(f"OK: '{cit}' matches something in References.")

# Check for table reference consistency
# Let's inspect the numbers in paragraphs vs tables.
# Let's write another helper to check table contents and find mismatches with text.
