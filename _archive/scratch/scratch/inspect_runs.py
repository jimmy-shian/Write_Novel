from docx import Document

doc = Document(r"C:\Users\user\Desktop\test_html\碩論\碩論vX.docx")

terms = [
    "base_no_rag",
    "base_with_rag",
    "bert_rag",
    "direct_code",
    "code_list_50",
    "index_rag",
    "Bert RAG 100",
    "Bert RAG Bonus",
    "Base Mix-Def",
    "PPL Base Mix-Def",
    "Indexrag Mix-Def",
    "Ppl Indexrag Mix-Def"
]

print("--- Paragraph Run Details ---")
for idx, p in enumerate(doc.paragraphs):
    text = p.text
    for term in terms:
        if term.lower() in text.lower():
            print(f"P {idx} (contains '{term}'):")
            print(f"  Full text: {text}")
            print(f"  Runs ({len(p.runs)}):")
            for r_idx, r in enumerate(p.runs):
                print(f"    Run {r_idx}: '{r.text}'")
            break

print("--- Table Cell Run Details ---")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text
            for term in terms:
                if term.lower() in text.lower():
                    print(f"Table {t_idx} R{r_idx}C{c_idx} (contains '{term}'):")
                    print(f"  Full text: {text}")
                    for p_idx, p in enumerate(cell.paragraphs):
                        print(f"    P {p_idx} Runs ({len(p.runs)}):")
                        for run_idx, r in enumerate(p.runs):
                            print(f"      Run {run_idx}: '{r.text}'")
                    break
