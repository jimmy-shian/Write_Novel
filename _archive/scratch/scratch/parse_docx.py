import os
import zipfile
import xml.etree.ElementTree as ET
from docx import Document

# XML Namespaces
namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

def extract_footnotes(docx_path):
    footnotes = []
    if not os.path.exists(docx_path):
        return footnotes
    try:
        with zipfile.ZipFile(docx_path) as z:
            if 'word/footnotes.xml' in z.namelist():
                xml_content = z.read('word/footnotes.xml')
                root = ET.fromstring(xml_content)
                for footnote in root.findall('.//w:footnote', namespaces):
                    f_id = footnote.get(f'{{{namespaces["w"]}}}id')
                    try:
                        if int(f_id) <= 0:
                            continue
                    except (ValueError, TypeError):
                        pass
                    
                    texts = [node.text for node in footnote.findall('.//w:t', namespaces) if node.text]
                    footnote_text = "".join(texts).strip()
                    if footnote_text:
                        footnotes.append((f_id, footnote_text))
    except Exception as e:
        print(f"Error extracting footnotes: {e}")
    return footnotes

def analyze_docx(docx_path, out_txt_path):
    doc = Document(docx_path)
    footnotes = extract_footnotes(docx_path)
    
    with open(out_txt_path, 'w', encoding='utf-8') as f:
        f.write(f"Document Analysis for: {docx_path}\n")
        f.write(f"Paragraphs: {len(doc.paragraphs)}\n")
        f.write(f"Tables: {len(doc.tables)}\n")
        f.write(f"Sections: {len(doc.sections)}\n")
        f.write(f"Footnotes found: {len(footnotes)}\n\n")
        
        f.write("=== FOOTNOTES ===\n")
        for f_id, text in footnotes:
            f.write(f"Footnote ID [{f_id}]: {text}\n")
        f.write("\n")
        
        f.write("=== PARAGRAPHS ===\n")
        for idx, p in enumerate(doc.paragraphs):
            f.write(f"Paragraph {idx}: {p.text}\n")
            
        f.write("\n=== TABLES ===\n")
        for idx, table in enumerate(doc.tables):
            f.write(f"\n--- Table {idx} ({len(table.rows)} rows, {len(table.columns)} cols) ---\n")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip().replace("\n", " "))
                f.write(f"Row {row_idx}: {' | '.join(row_text)}\n")

if __name__ == '__main__':
    analyze_docx(
        r"C:\Users\user\Desktop\test_html\碩論\碩論vX.docx",
        r"C:\Users\user\Desktop\test_html\Write_Novel\scratch\doc_dump.txt"
    )
    print("Dump completed.")
