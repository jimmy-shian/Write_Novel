import re
from docx import Document

doc = Document(r"C:\Users\user\Desktop\test_html\碩論\碩論vX.docx")

mapping = {
    "base_no_rag": "Base",
    "base_with_rag": "RAG",
    "bert_rag": "RefineRAG",
    "direct_code": "CodeRAG",
    "code_list_50": "Top50RAG",
    "index_rag": "IndexRAG",
    "Bert RAG 100": "AC-RAG",
    "Bert RAG Bonus": "TeachRAG",
    "Base Mix-Def": "DefBase",
    "PPL Base Mix-Def": "PPL-DefBase",
    "Indexrag Mix-Def": "DefIndex",
    "Ppl Indexrag Mix-Def": "PPL-DefIndex"
}

def check_all():
    # We want to search for any case-insensitive or minor variant matching of keys
    # Let's compile a list of regexes for each key.
    # Note: we need to order keys by length descending to match longer ones first (e.g. "Bert RAG 100" before "Bert RAG").
    sorted_keys = sorted(mapping.keys(), key=len, reverse=True)
    
    # We will search paragraph by paragraph and table by table.
    print("=== Paragraph Matches ===")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        if not text.strip():
            continue
        # Find which keys match in text (case-insensitive)
        matched_keys = []
        for key in sorted_keys:
            # Escape space and underscore to match both or either
            pattern = key.replace(" ", r"[\s_]*").replace("_", r"[\s_]*")
            if re.search(pattern, text, re.IGNORECASE):
                # Print match
                matched_keys.append(key)
        if matched_keys:
            print(f"P {idx} (matched {matched_keys}): {text[:150]}")

    print("\n=== Table Matches ===")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                if not text.strip():
                    continue
                matched_keys = []
                for key in sorted_keys:
                    pattern = key.replace(" ", r"[\s_]*").replace("_", r"[\s_]*")
                    if re.search(pattern, text, re.IGNORECASE):
                        matched_keys.append(key)
                if matched_keys:
                    print(f"Table {t_idx} R{r_idx}C{c_idx} (matched {matched_keys}): {text}")

if __name__ == '__main__':
    check_all()
