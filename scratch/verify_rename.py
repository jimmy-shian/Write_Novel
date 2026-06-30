import re
from docx import Document

mapping = {
    "Ppl Indexrag Mix-Def": "PPL-DefIndex",
    "Indexrag Mix-Def": "DefIndex",
    "PPL Base Mix-Def": "PPL-DefBase",
    "Base Mix-Def": "DefBase",
    "Bert RAG Bonus": "TeachRAG",
    "Bert RAG 100": "AC-RAG",
    "bert_rag": "RefineRAG",
    "base_no_rag": "Base",
    "base_with_rag": "RAG",
    "direct_code": "CodeRAG",
    "code_list_50": "Top50RAG",
    "index_rag": "IndexRAG",
}

def verify():
    docx_path = r"C:\Users\user\Desktop\test_html\碩論\碩論vX_命名更新.docx"
    doc = Document(docx_path)
    
    # Check for original terms in paragraphs
    original_found = 0
    new_found = {k: 0 for k in mapping.values()}
    
    # Standard check list
    sorted_keys = sorted(mapping.keys(), key=len, reverse=True)
    
    print("=== Scanning Updated Paragraphs for Original Terms ===")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        for key in sorted_keys:
            pattern = key.replace(" ", r"[\s_]*").replace("_", r"[\s_]*")
            if re.search(pattern, text, re.IGNORECASE):
                # Make sure it's not matching the updated ones (like IndexRAG matching index_rag pattern)
                # But wait, index_rag pattern is `index[_\s-]?rag`. This matches IndexRAG.
                # So we must check if it is exactly the new term
                if "IndexRAG" in text and key == "index_rag" and not re.search(r'\bindex_rag\b', text, re.IGNORECASE):
                    # It matched IndexRAG, which is the new term. That's fine!
                    continue
                if "AC-RAG" in text and key == "bert_rag":
                    # AC-RAG contains RAG, wait, bert_rag pattern is `bert[_\s-]?rag` which won't match AC-RAG.
                    pass
                print(f"P {idx} still contains a match for '{key}': {text}")
                original_found += 1
                
        # Count new terms
        for new_val in new_found.keys():
            if new_val in text:
                new_found[new_val] += 1

    print("\n=== Scanning Updated Tables for Original Terms ===")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                for key in sorted_keys:
                    pattern = key.replace(" ", r"[\s_]*").replace("_", r"[\s_]*")
                    if re.search(pattern, text, re.IGNORECASE):
                        if "IndexRAG" in text and key == "index_rag" and not re.search(r'\bindex_rag\b', text, re.IGNORECASE):
                            continue
                        print(f"Table {t_idx} R{r_idx}C{c_idx} still contains '{key}': {text}")
                        original_found += 1
                        
                for new_val in new_found.keys():
                    if new_val in text:
                        new_found[new_val] += 1

    print("\n=== Summary of New Terms Found ===")
    for k, v in new_found.items():
        print(f"New term '{k}': found {v} times.")
        
    print(f"\nVerification status: {original_found} original terms remaining.")

if __name__ == '__main__':
    verify()
