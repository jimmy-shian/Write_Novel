import re
from docx import Document

mapping = {
    "Ppl Indexrag Mix-Def": "PPL-DefIndex",
    "Indexrag Mix-Def": "DefIndex",
    "PPL Base Mix-Def": "PPL-DefBase",
    "Base Mix-Def": "DefBase",
    "Bert RAG Bonus": "TeachRAG",
    "Bert RAG 100": "AC-RAG",
    "bert_rag": "RefineRAG",
    "base_no_rag": "Base",
    "base_with_rag": "RAG",
    "direct_code": "CodeRAG",
    "code_list_50": "Top50RAG",
    "index_rag": "IndexRAG",
}

# Patterns mapping to target replacement values
patterns = [
    (r'\bppl[_\s-]index[_\s-]?rag[_\s-]mix[_\s-]def\b', "PPL-DefIndex"),
    (r'\bindex[_\s-]?rag[_\s-]mix[_\s-]def\b', "DefIndex"),
    (r'\bppl[_\s-]base[_\s-]mix[_\s-]def\b', "PPL-DefBase"),
    (r'\bbase[_\s-]mix[_\s-]def\b', "DefBase"),
    (r'\bbert[_\s-]?rag[_\s-]bonus\b', "TeachRAG"),
    (r'\bbert[_\s-]?rag[_\s-]100\b', "AC-RAG"),
    # Matches bert_rag, Bert RAG, BertRAG
    (r'\bbert[_\s-]?rag\b', "RefineRAG"),
    (r'\bbase[_\s-]no[_\s-]rag\b', "Base"),
    (r'\bbase[_\s-]with[_\s-]rag\b', "RAG"),
    (r'\bdirect[_\s-]?code\b', "CodeRAG"),
    (r'\bcode[_\s-]?list[_\s-]?50\b', "Top50RAG"),
    (r'\bindex[_\s-]?rag\b', "IndexRAG"),
]

def replace_text(text):
    original = text
    replaced = False
    for pattern, val in patterns:
        new_text, count = re.subn(pattern, val, text, flags=re.IGNORECASE)
        if count > 0:
            text = new_text
            replaced = True
    return text, replaced

def process_paragraph(p):
    replaced_any = False
    original_text = p.text
    # First, try replacing within individual runs to preserve formatting as much as possible
    for run in p.runs:
        new_run_text, run_replaced = replace_text(run.text)
        if run_replaced:
            run.text = new_run_text
            replaced_any = True
            
    # Always check if the full paragraph text still contains any target pattern.
    full_text, text_replaced = replace_text(p.text)
    if text_replaced:
        # Check if the text actually changed from the current state of p.text
        # (which might have already been partially updated by run replacements)
        if p.text != full_text:
            print(f"Fallback warning: Reconstructing paragraph text due to split runs in paragraph: {original_text[:100]}...")
            p.text = full_text
            replaced_any = True
        
    return replaced_any

def main():
    docx_path = r"C:\Users\user\Desktop\test_html\碩論\碩論vX.docx"
    output_path = r"C:\Users\user\Desktop\test_html\碩論\碩論vX_命名更新.docx"
    
    doc = Document(docx_path)
    
    paragraph_replacements = 0
    table_replacements = 0
    
    # Process regular paragraphs
    for idx, p in enumerate(doc.paragraphs):
        if process_paragraph(p):
            paragraph_replacements += 1
            
    # Process tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    if process_paragraph(p):
                        table_replacements += 1
                        
    # Process headers and footers in all sections
    header_footer_replacements = 0
    for s_idx, section in enumerate(doc.sections):
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for p in header.paragraphs:
                    if process_paragraph(p):
                        header_footer_replacements += 1
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if process_paragraph(p):
                                    header_footer_replacements += 1
                                    
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for p in footer.paragraphs:
                    if process_paragraph(p):
                        header_footer_replacements += 1
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if process_paragraph(p):
                                    header_footer_replacements += 1

    print(f"Completed replacements:")
    print(f"Paragraphs updated: {paragraph_replacements}")
    print(f"Table cells/paragraphs updated: {table_replacements}")
    print(f"Headers/footers updated: {header_footer_replacements}")
    
    doc.save(output_path)
    print(f"Saved updated file to: {output_path}")

if __name__ == '__main__':
    main()
